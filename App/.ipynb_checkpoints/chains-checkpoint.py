import os
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()


def add_hyperlink(paragraph, text, url):
    """A function to add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", r_id
    )
    new_run = OxmlElement("w:r")
    r_text = OxmlElement("w:t")
    r_text.text = text
    new_run.append(r_text)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)


class Chain:
    def __init__(self):
        self.llm = ChatGroq(
            temperature=0,
            groq_api_key=os.getenv("GROQ_API_KEY"),
            model_name="llama-3.1-70b-versatile",
        )

    def extract_jobs(self, cleaned_text):
        prompt_extract = PromptTemplate.from_template(
            """
            ### SCRAPED TEXT FROM WEBSITE:
            {page_data}
            ### INSTRUCTION:
            The scraped text is from the career's page of a website.
            Your job is to extract the job postings and return them in JSON format containing the following keys: `role`, `experience`, `skills` and `description`.
            Only return the valid JSON.
            ### VALID JSON (NO PREAMBLE):
            """
        )
        chain_extract = prompt_extract | self.llm
        res = chain_extract.invoke(input={"page_data": cleaned_text})
        try:
            json_parser = JsonOutputParser()
            res = json_parser.parse(res.content)
        except OutputParserException:
            raise OutputParserException("Context too big. Unable to parse jobs.")
        return res if isinstance(res, list) else [res]

    def write_mail(self, job, links):
        prompt_email = PromptTemplate.from_template(
            """
            ### JOB DESCRIPTION:
            {job_description}

            ### INSTRUCTION:
            You are Pradhum Niroula, a business analytics student at Oakland University.
            Your job is to write a cold email to the HR regarding the job mentioned above describing my capability
            in fulfilling their needs.
            Also add the most relevant ones from the following links to showcase Pradhum's portfolio: {link_list}
            Remember you are Pradhum Niroula, Master of Science in Business Analytics at Oakland University. 
            Do not provide a preamble.
            ### EMAIL (NO PREAMBLE):

            """
        )
        chain_email = prompt_email | self.llm
        res = chain_email.invoke({"job_description": str(job), "link_list": links})
        return res.content

    def write_content(self, resume_content, job_description, links, content_type="cover_letter"):
        if content_type == "cover_letter":
            return self.write_cover_letter(resume_content, job_description, links)
        else:
            raise ValueError("Unsupported content type")

    def write_cover_letter(self, resume_content, job_description, links):
        prompt_cover_letter = PromptTemplate.from_template(
            """
            ### INSTRUCTION:
             You are a job seeker applying for the position described in the job posting below. Your goal is to write a professional cover letter.
             Use the provided resume and the job description to highlight your relevant experience, skills, and why you're a great fit for the company and the role.
             Include the most relevant items from the following portfolio links: {link_list}
             ### RESUME:
             {resume_content}
             ### JOB DESCRIPTION:
             {job_description}
             ### COVER LETTER (NO PREAMBLE):
            """
        )
        chain_cover_letter = prompt_cover_letter | self.llm
        res = chain_cover_letter.invoke(
            {
                "resume_content": resume_content,
                "job_description": job_description,
                "link_list": links,
            }
        )
        return res.content

    def save_cover_letter(self, content, filename="Cover_Letter.docx"):
        doc = Document()

        # Add name as the header
        heading = doc.add_paragraph()
        heading.add_run("Pradhum Niroula").bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add clickable contact details
        contact = doc.add_paragraph()
        add_hyperlink(contact, "+1 (947) 276-3480", "tel:+19472763480")
        contact.add_run(" || ")
        add_hyperlink(contact, "pradesgniroula55@gmail.com", "mailto:pradesgniroula55@gmail.com")
        contact.add_run(" || Auburn Hills, MI")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add top horizontal line
        doc.add_paragraph()
        line = doc.add_paragraph()
        line.add_run("─────────────────────────────────────────────────────────").bold = True
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date
        doc.add_paragraph("\nSeptember 28, 2024")

        # Recipient details
        recipient = doc.add_paragraph()
        recipient.add_run("Mr. Jeff Bennett\n").bold = True
        recipient.add_run("Talent Acquisition Senior Manager, Deloitte LLP\n")
        recipient.add_run("1001 Woodward,\n")
        recipient.add_run("MI, 48226-1904\n")

        # Salutation
        doc.add_paragraph("\nDear Sir,\n")

        # Body content
        paragraphs = content.split("\n\n")
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph).style.font.size = Pt(12)

        # Closing statement
        doc.add_paragraph("\nSincerely,")
        doc.add_paragraph("Pradhum Niroula")

        # Save file
        doc.save(filename)
        return filename
