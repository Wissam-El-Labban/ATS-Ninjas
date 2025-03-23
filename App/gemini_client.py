import os
import re
import json
from dotenv import load_dotenv

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.exceptions import OutputParserException

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

load_dotenv()

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", r_id
    )
    new_run = OxmlElement("w:r")
    r_text = OxmlElement("w:t")
    r_text.text = text
    new_run.append(r_text)
    hyperlink.append(new_run)
    paragraph._element.append(hyperlink)

class GeminiClient:
    def __init__(self):
        self.llm = ChatGoogleGenerativeAI(
            google_api_key=os.getenv("GOOGLE_API_KEY"),
            temperature=0,
            model="models/gemini-2.0-pro-exp-02-05"
        )

    def extract_jobs(self, cleaned_text):
        prompt_extract = PromptTemplate.from_template(
            """
            ### SCRAPED TEXT FROM WEBSITE:
            {page_data}

            ### INSTRUCTION:
            The scraped text is from a careers page.
            Extract job postings and return them in valid JSON with:
            role, experience, skills, description.
            """
        )
        chain_extract = prompt_extract | self.llm
        res = chain_extract.invoke({"page_data": cleaned_text})
        raw = res.content.strip()

        # Remove numeric keys + trailing commas, fix quotes
        raw = re.sub(r"\n?\s*\d+\s*:", "", raw)
        raw = re.sub(r",(\s*[\]}])", r"\1", raw)
        raw = raw.replace("“", "\"").replace("”", "\"")
        raw = raw.replace("‘", "'").replace("’", "'")

        try:
            json_parser = JsonOutputParser()
            parsed = json_parser.parse(raw)
            return parsed if isinstance(parsed, list) else [parsed]
        except OutputParserException as e:
            raise OutputParserException(f"❌ Failed to parse JSON:\n\n{raw}\n\nError: {e}")

    def write_mail(self, job_description, links):
        prompt_email = PromptTemplate.from_template(
            """
            ### JOB DESCRIPTION:
            {job_desc}

            ### TASK:
            Write a cold email from Pradhum Niroula about the above job.
            Reference these links: {link_list}.
            Return only the email text.
            """
        )
        chain_email = prompt_email | self.llm
        res = chain_email.invoke({"job_desc": str(job_description), "link_list": links})
        return res.content

    def write_cover_letter(self, resume, job_description, links):
        prompt_cover = PromptTemplate.from_template(
            """
            ### INSTRUCTION:
            You are Pradhum Niroula applying for a job.
            - Resume: {resume_data}
            - Description: {job_desc}
            - Links: {link_list}
            Write a cover letter with no extra text.
            """
        )
        chain_cover = prompt_cover | self.llm
        res = chain_cover.invoke({
            "resume_data": resume,
            "job_desc": job_description,
            "link_list": links
        })
        return res.content

    def save_cover_letter(self, content, filename="Cover_Letter.docx"):
        doc = Document()

        heading = doc.add_paragraph()
        heading.add_run("Pradhum Niroula").bold = True
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = doc.add_paragraph()
        add_hyperlink(contact, "+1 (947) 276-3480", "tel:+19472763480")
        contact.add_run(" || ")
        add_hyperlink(contact, "pradesgniroula55@gmail.com", "mailto:pradesgniroula55@gmail.com")
        contact.add_run(" || Auburn Hills, MI")
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        line = doc.add_paragraph()
        line.add_run("─────────────────────────────────────────────────────────").bold = True
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("\nSeptember 28, 2024")

        recipient = doc.add_paragraph()
        recipient.add_run("Mr. Jeff Bennett\n").bold = True
        recipient.add_run("Talent Acquisition Senior Manager, Deloitte LLP\n")
        recipient.add_run("1001 Woodward,\n")
        recipient.add_run("MI, 48226-1904\n")

        doc.add_paragraph("\nDear Sir,\n")

        paragraphs = content.split("\n\n")
        for paragraph in paragraphs:
            doc.add_paragraph(paragraph).style.font.size = Pt(12)

        doc.add_paragraph("\nSincerely,")
        doc.add_paragraph("Pradhum Niroula")

        doc.save(filename)
        return filename
